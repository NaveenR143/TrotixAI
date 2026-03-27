from __future__ import annotations

import abc
import io
import logging
import csv
import os
import tempfile
import uuid
import re
from collections import defaultdict
from fuzzywuzzy import fuzz

from .errors import ParsingError

LOGGER = logging.getLogger("resume_processor")

try:
    import pdfplumber  # type: ignore
except Exception:  # pragma: no cover
    pdfplumber = None

try:
    import fitz  # PyMuPDF  # type: ignore
except Exception:  # pragma: no cover
    fitz = None

try:
    from docx import Document as DocxDocument  # type: ignore
except Exception:  # pragma: no cover
    DocxDocument = None  # type: ignore

try:
    import textract  # type: ignore
except Exception:  # pragma: no cover
    textract = None


class ResumeParserStrategy(abc.ABC):
    """Strategy interface for file-type-specific text extraction."""

    @abc.abstractmethod
    def parse_text(self, file_bytes: bytes) -> str:
        raise NotImplementedError


class ParsingError(Exception):
    pass


class PdfResumeParser(ResumeParserStrategy):
    """Advanced resume parser with table-aware + semantic grouping."""

    BLOCK_SEPARATOR = "\n<<<END_OF_BLOCK_{id}>>>\n"

    SECTION_KEYWORDS = {
        "summary": ["summary", "profile summary", "professional summary", "about", "profile"],
        "objective": ["career objective", "objective", "career goal", "goal"],
        "skills": ["skills", "core competencies", "competencies", "technical skills", "expertise", "technical expertise"],
        "education": ["education", "academics", "qualification", "qualifications", "educational", "schooling"],
        "certifications": ["certification", "certifications", "courses", "training"],
        "projects": ["project", "projects", "academic project", "professional project"],
        "experience": ["experience", "work experience", "employment", "professional experience", "work"],
        "achievements": ["achievements", "awards", "awards and recognition"],
        "strengths": ["strength", "areas of strength", "key strengths"],
        "personal": ["personal", "personal details", "personal information", "dob", "date of birth"],
        "declaration": ["declaration"],
        "references": ["references", "reference"],
        "languages": ["languages", "language proficiency"],
        "interests": ["interests", "hobbies"],
    }

    # ------------------ HELPERS ------------------

    def _format_block(self, section: str, content: str) -> str:
        try:
            if not section or not isinstance(section, str):
                raise ValueError(f"Invalid section: {section}")
            if not isinstance(content, str):
                raise ValueError(
                    f"Content must be string, got {type(content)}")

            block_id = uuid.uuid4().hex[:8]
            header = f"\n### {section.upper()} ###\n"
            formatted = header + content.strip() + self.BLOCK_SEPARATOR.format(id=block_id)
            return formatted
        except Exception as exc:
            print(f"Error formatting block for section '{section}': {exc}")
            LOGGER.error(
                f"Error formatting block for section '{section}': {exc}")
            raise

    def _normalize_text(self, text: str) -> str:
        """Normalize text by removing excessive internal whitespace but preserving paragraphs."""
        try:
            if not text or not isinstance(text, str):
                return ""

            # Split by double newlines (paragraph breaks) first
            lines = text.split("\n")
            normalized_lines = []

            for line in lines:
                # On each line, replace excessive internal spaces with single space
                # Multiple spaces/tabs -> single space
                normalized = re.sub(r'[ \t]{2,}', ' ', line)
                normalized = normalized.strip()
                if normalized:  # Only keep non-empty lines
                    normalized_lines.append(normalized)

            # Rejoin with newlines
            return "\n".join(normalized_lines)
        except Exception as exc:
            LOGGER.debug(f"Error normalizing text: {exc}")
            return text.strip() if text else ""

    def _clean_line(self, line: str) -> str:
        try:
            if not isinstance(line, str):
                return str(line).strip()

            # First normalize excessive whitespace
            line = self._normalize_text(line)

            # Remove leading special characters and bullets
            cleaned = re.sub(r'^[\W_]+', '', line)
            return cleaned.strip()
        except Exception as exc:
            print(f"Error cleaning line: {exc}")
            LOGGER.warning(f"Error cleaning line: {exc}")
            return str(line).strip()

    def _is_all_caps(self, text: str) -> bool:
        try:
            if not text or not isinstance(text, str):
                return False
            return text.isupper() and len(text.split()) <= 6
        except Exception as exc:
            print(f"Error checking all caps: {exc}")
            LOGGER.warning(f"Error checking all caps: {exc}")
            return False

    def _is_potential_heading(self, line, is_bold):
        try:
            if not line or not isinstance(line, str):
                return False

            words = line.split()
            if len(words) == 0 or len(words) > 8:  # Too many words = not a header
                return False

            # Avoid table-like content and long lines
            # Avoid "key: value" lines
            if "|" in line or ":" in line.rstrip(':'):
                return False

            formatting = self._analyze_text_formatting(line)

            # Only consider ALL CAPS + short as potential header
            # Don't use other formatting cues to reduce false positives
            if formatting["is_uppercase"] and formatting["is_short"] and formatting["word_count"] <= 4:
                return True

            # Only accept bold if explicitly provided (not commonly available from pdfplumber)
            if is_bold and formatting["word_count"] <= 4:
                return True

            return False
        except Exception as exc:
            LOGGER.debug(f"Error checking potential heading: {exc}")
            return False

    def _is_heading_fuzzy(self, line: str):
        try:
            if not line or not isinstance(line, str):
                return None

            line_clean = line.lower().strip()
            # Headers shouldn't be too long
            if not line_clean or len(line_clean) > 100:
                return None

            best_match = None
            best_score = 0

            for section, keywords in self.SECTION_KEYWORDS.items():
                for kw in keywords:
                    try:
                        score = fuzz.partial_ratio(line_clean, kw)
                        # Be more strict with scoring - require 70+ for matches
                        if score > best_score and score >= 70:
                            best_score = score
                            best_match = section
                    except Exception as kw_exc:
                        LOGGER.debug(f"Error fuzzy matching '{kw}': {kw_exc}")
                        continue

            # Higher threshold for fuzzy matching to reduce false positives
            if best_score >= 75:
                return best_match
            return None
        except Exception as exc:
            print(f"Error in fuzzy heading detection: {exc}")
            LOGGER.error(f"Error in fuzzy heading detection: {exc}")
            return None

    def _analyze_text_formatting(self, line: str) -> dict:
        """Analyze text formatting cues (uppercase, spacing, length, symbols)."""
        try:
            if not line or not isinstance(line, str):
                return {"is_uppercase": False, "is_short": False, "has_multiple_spaces": False}

            stripped = line.strip()
            if not stripped:
                return {"is_uppercase": False, "is_short": False, "has_multiple_spaces": False}

            return {
                "is_uppercase": stripped.isupper() and len(stripped.split()) > 0,
                "is_short": len(stripped.split()) <= 6,
                "has_multiple_spaces": "  " in line,  # Extra spacing can indicate headers
                "line_length": len(stripped),
                "word_count": len(stripped.split()),
                "has_colon": ":" in stripped,
                "starts_with_number_or_bullet": bool(re.match(r'^[0-9+\-*\u2022\u00b7]', stripped))
            }
        except Exception as exc:
            LOGGER.debug(f"Error analyzing text formatting: {exc}")
            return {"is_uppercase": False, "is_short": False, "has_multiple_spaces": False}

    # ------------------ TABLE LOGIC ------------------

    def _format_table_as_tsv(self, table, table_type="unknown"):
        """Format table as tab-separated values with header row."""
        try:
            if not table or len(table) == 0:
                return []

            formatted_rows = []

            # Add header row
            if table:
                header = table[0]
                if header:
                    # Clean header cells
                    clean_header = [
                        str(cell).strip() if cell else "" for cell in header]
                    header_row = "\t".join(clean_header)
                    formatted_rows.append(header_row)

            # Add data rows
            for row_idx, row in enumerate(table[1:], 1):
                try:
                    if not row:
                        continue

                    # Clean cell values
                    clean_cells = [
                        str(cell).strip() if cell else "" for cell in row]
                    data_row = "\t".join(clean_cells)

                    if data_row.strip():  # Only add non-empty rows
                        formatted_rows.append(data_row)
                except Exception as row_exc:
                    LOGGER.debug(
                        f"Error formatting table row {row_idx}: {row_exc}")
                    continue

            return formatted_rows
        except Exception as exc:
            print(f"Error formatting table as TSV: {exc}")
            LOGGER.error(f"Error formatting table as TSV: {exc}")
            return []

    def _parse_education_table(self, table):
        try:
            if not table or len(table) < 2:
                return []

            # Return as TSV format
            return self._format_table_as_tsv(table, "education")
        except Exception as exc:
            print(f"Error parsing education table: {exc}")
            LOGGER.error(f"Error parsing education table: {exc}")
            return []

    def _parse_generic_table(self, table):
        try:
            if not table or len(table) < 2:
                return []

            # Return as TSV format
            return self._format_table_as_tsv(table, "generic")
        except Exception as exc:
            print(f"Error parsing generic table: {exc}")
            LOGGER.error(f"Error parsing generic table: {exc}")
            return []

    def _parse_skills_table(self, table):
        try:
            if not table:
                return []

            # For skills, we might want a simpler format
            # Return as TSV format
            return self._format_table_as_tsv(table, "skills")
        except Exception as exc:
            print(f"Error parsing skills table: {exc}")
            LOGGER.error(f"Error parsing skills table: {exc}")
            return []

    def _parse_experience_table(self, table):
        try:
            if not table or len(table) < 2:
                return []

            # Return as TSV format
            return self._format_table_as_tsv(table, "experience")
        except Exception as exc:
            print(f"Error parsing experience table: {exc}")
            LOGGER.error(f"Error parsing experience table: {exc}")
            return []

    def _parse_projects_table(self, table):
        try:
            if not table:
                return []

            # Return as TSV format
            return self._format_table_as_tsv(table, "projects")
        except Exception as exc:
            print(f"Error parsing projects table: {exc}")
            LOGGER.error(f"Error parsing projects table: {exc}")
            return []

    def _classify_table(self, table):
        try:
            if not table or not isinstance(table, list):
                return "unknown"

            if len(table) == 0 or not table[0]:
                return "unknown"

            # Extract and validate headers
            headers = " ".join([str(c).lower() for c in table[0] if c])
            if not headers:
                return "unknown"

            # Education keywords
            if any(k in headers for k in ["university", "college", "degree", "qualification", "board"]):
                return "education"

            # Experience keywords
            if any(k in headers for k in ["company", "role", "designation", "duration"]):
                return "experience"

            # Projects keywords
            if any(k in headers for k in ["project", "description", "technology"]):
                return "projects"

            # Skills keywords
            if any(k in headers for k in ["skill", "technology", "tools"]):
                return "skills"

            # Fallback: scan all cells
            try:
                flat = " ".join(str(cell).lower()
                                for row in table for cell in row if cell)
                if not flat:
                    return "unknown"

                if "sql" in flat or "python" in flat or "java" in flat:
                    return "skills"

                if "bachelor" in flat or "university" in flat or "college" in flat:
                    return "education"
            except Exception as scan_exc:
                LOGGER.debug(f"Error scanning table cells: {scan_exc}")

            return "unknown"
        except Exception as exc:
            print(f"Error classifying table: {exc}")
            LOGGER.error(f"Error classifying table: {exc}")
            return "unknown"

    def _parse_projects_table(self, table):
        try:
            if not table:
                return []

            # Return as TSV format
            return self._format_table_as_tsv(table, "projects")
        except Exception as exc:
            print(f"Error parsing projects table: {exc}")
            LOGGER.error(f"Error parsing projects table: {exc}")
            return []

    def _process_tables(self, tables, section_map):
        try:
            for table_idx, table in enumerate(tables):
                try:
                    if not table or len(table) < 2:
                        LOGGER.debug(f"Skipping invalid table {table_idx}")
                        continue
                    
                    table_type = self._classify_table(table)
                    
                    # Format table as TSV
                    formatted_rows = self._format_table_as_tsv(table, table_type)
                    
                    if formatted_rows:
                        # Add formatted table to appropriate section
                        if table_type != "unknown":
                            section_map[table_type].extend(formatted_rows)
                            LOGGER.info(f"Added {len(formatted_rows)} rows to {table_type} table")
                        else:
                            # Add unknown tables to education or general
                            section_map["education"].extend(formatted_rows)
                            LOGGER.debug(f"Added unknown table ({len(formatted_rows)} rows) to education")
                except Exception as table_exc:
                    print(f"Error processing table {table_idx}: {table_exc}")
                    LOGGER.warning(f"Error processing table {table_idx}: {table_exc}")
                    continue
        except Exception as exc:
            print(f"Error in _process_tables: {exc}")
            LOGGER.error(f"Error in _process_tables: {exc}")
            raise

    # ------------------ EXTRACTION ------------------

    def _get_table_bounding_boxes(self, page, tables):
        """Get bounding boxes of all tables on a page."""
        try:
            table_bboxes = []
            if tables:
                for table in tables:
                    try:
                        # Get table cells bounding box
                        if hasattr(table, 'bbox') and table.bbox:
                            table_bboxes.append(table.bbox)
                        else:
                            # Fallback: try to extract coordinates from table settings
                            pass
                    except Exception as e:
                        LOGGER.debug(f"Could not get table bbox: {e}")
            return table_bboxes
        except Exception as exc:
            LOGGER.debug(f"Error getting table bounding boxes: {exc}")
            return []

    def _text_in_table_region(self, char, table_bboxes):
        """Check if a character position is within any table region."""
        try:
            if not char or not hasattr(char, 'x0') or not hasattr(char, 'x1'):
                return False
            
            char_bbox = (char.x0, char.top, char.x1, char.bottom)
            
            for table_bbox in table_bboxes:
                # Check if character overlaps with table
                if (char_bbox[0] < table_bbox[2] and char_bbox[2] > table_bbox[0] and
                    char_bbox[1] < table_bbox[3] and char_bbox[3] > table_bbox[1]):
                    return True
            return False
        except Exception:
            return False

    def _extract_text_excluding_tables(self, page, tables):
        """Extract text from page while excluding table regions."""
        try:
            # Try to extract text excluding tables using layout mode
            text = page.extract_text(layout=False) or ""
            
            # Alternative: use table detection settings to exclude table regions
            # For now, we'll use a simple heuristic approach
            if not text:
                text = page.extract_text() or ""
            
            return text
        except Exception as e:
            LOGGER.debug(f"Error extracting text excluding tables: {e}")
            return ""

    def _extract_with_pdfplumber(self, file_bytes: bytes):
        """Extract text and tables from PDF, preserving all content separately."""
        try:
            import pdfplumber

            text_content = []
            tables_content = []
            extracted_text_blocks = []
            unclassified_text = []

            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                page_count = len(pdf.pages)
                LOGGER.info(f"Processing PDF with {page_count} pages")

                for page_idx, page in enumerate(pdf.pages):
                    try:
                        # Extract tables FIRST to get their coordinates
                        tables_on_page = []
                        try:
                            tables = page.extract_tables()
                            if tables:
                                for table_idx, table in enumerate(tables):
                                    # Only keep meaningful tables
                                    if table and len(table) > 1:
                                        tables_on_page.append(table)
                                        tables_content.append({
                                            "page": page_idx + 1,
                                            "table_index": table_idx,
                                            "data": table
                                        })
                                        LOGGER.debug(
                                            f"Page {page_idx + 1}: Extracted table {table_idx} with {len(table)} rows")
                        except Exception as table_exc:
                            LOGGER.debug(f"Table extraction error on page {page_idx + 1}: {table_exc}")

                        # Extract text now (which may include table text)
                        text = page.extract_text() or ""

                        if text:
                            # Remove lines that look like table content (connected directly to a table)
                            # For now, keep the text as-is and let the section detection filters it
                            text_content.append(text)
                            extracted_text_blocks.append({
                                "page": page_idx + 1,
                                "text": text,
                                "length": len(text)
                            })
                            LOGGER.debug(
                                f"Page {page_idx + 1}: Extracted {len(text)} chars of text")
                        else:
                            # Try alternative extraction
                            try:
                                alt_text = page.extract_text(layout=True) or ""
                                if alt_text:
                                    text_content.append(alt_text)
                                    extracted_text_blocks.append({
                                        "page": page_idx + 1,
                                        "text": alt_text,
                                        "length": len(alt_text),
                                        "method": "layout"
                                    })
                                    LOGGER.debug(f"Page {page_idx + 1}: Layout extraction retrieved {len(alt_text)} chars")
                            except Exception as alt_exc:
                                LOGGER.debug(f"Layout fallback failed: {alt_exc}")

                    except Exception as page_exc:
                        print(f"Error extracting from page {page_idx + 1}: {page_exc}")
                        LOGGER.error(f"Error extracting from page {page_idx + 1}: {page_exc}")
                        unclassified_text.append({"page": page_idx + 1, "error": str(page_exc)})

            # Join all text
            full_text = "\n".join(text_content)
            LOGGER.info(
                f"Total extracted text: {len(full_text)} chars from {len(extracted_text_blocks)} blocks, {len(tables_content)} tables")

            if len(full_text.strip()) == 0:
                print("WARNING: No text content extracted from PDF. File may be image-based.")
                LOGGER.warning("No text content extracted from PDF")

            # Return text and tables separately
            return full_text, [t["data"] for t in tables_content]
        except Exception as exc:
            print(f"Error in PDF extraction: {exc}")
            LOGGER.error(f"Error in PDF extraction: {exc}")
            raise

    def _is_likely_table_line(self, line: str) -> bool:
        """Detect if a line appears to be extracted from table content."""
        try:
            if not line or not isinstance(line, str):
                return False
            
            stripped = line.strip()
            if not stripped or len(stripped) < 3:
                return False
            
            # Look for table header patterns (all caps keywords with few words)
            table_headers = [
                "QUALIFICATION", "UNIVERSITY", "BOARD", "INSTITUTION", "YEAR",
                "COMPLETION", "PERCENTAGE", "COMPANY", "DESIGNATION", "DURATION",
                "SKILL", "TECHNOLOGY", "PROFICIENCY", "PROJECT", "DESCRIPTION"
            ]
            
            # If line has multiple table keywords, it's likely a table header
            header_keywords = sum(1 for keyword in table_headers if keyword in stripped.upper())
            if header_keywords >= 2:
                return True
            
            words = stripped.split()
            
            if len(words) < 3:
                return False
            
            # Count different types of tokens
            year_like = sum(1 for w in words if w.isdigit() and 1900 <= int(w) <= 2050)
            percent_like = sum(1 for w in words if w.replace('.', '', 1).isdigit() and '.' in w)
            
            # SPECIFIC: Degree keywords pattern with years AND percentages (education table)
            degree_keywords = ["bachelor", "intermediate", "10th", "masters", "diploma", "bca", "btech", "mca", "mtech", "phd"]
            has_degree = any(kw in stripped.lower() for kw in degree_keywords)
            
            # Education table pattern: degree + org names + year + percentage
            if has_degree and year_like > 0 and percent_like > 0:
                return True
            
            # Company/experience table pattern: company name + role + year
            company_keywords = ["company", "pvt", "ltd", "inc", "llc", ".com"]
            has_company = any(kw.lower() in stripped.lower() for kw in company_keywords)
            if has_company and year_like > 0:
                return True
            
            # Avoid false positives for normal sentences with years
            # e.g., "2015 Certification in Oracle..."
            if any(ending in stripped for ending in ['.', '!', '?']):
                # Has sentence ending, likely normal text despite year
                return False
            
            return False
        except Exception as exc:
            LOGGER.debug(f"Error detecting table line: {exc}")
            return False

    # ------------------ MAIN ------------------

    def parse_text(self, file_bytes: bytes) -> str:
        """Parse PDF ensuring all text is captured and organized into logical sections."""
        try:
            text, tables = self._extract_with_pdfplumber(file_bytes)

            if not text:
                raise ParsingError("Unable to extract PDF text.")

            # Normalize extracted text to remove excessive whitespace
            text = self._normalize_text(text)

            # Split into lines for processing
            raw_lines = [line for line in text.split("\n") if line.strip()]

            # Log initial stats
            total_lines = len(raw_lines)
            LOGGER.info(f"Total lines after normalization: {total_lines}")

            section_map = defaultdict(list)
            current_section = "general"
            buffer = []  # Buffer for accumulating content lines

            for line_idx, line in enumerate(raw_lines):
                try:
                    cleaned_line = self._clean_line(line)

                    if not cleaned_line:
                        continue

                    # ---- Skip lines that are likely from table extraction ----
                    if self._is_likely_table_line(cleaned_line):
                        LOGGER.debug(f"Skipping table line: {cleaned_line[:50]}")
                        continue

                    # ---- detect heading with multiple cues ----
                    detected_section = self._is_heading_fuzzy(cleaned_line)

                    if not detected_section and self._is_potential_heading(cleaned_line, False):
                        detected_section = self._is_heading_fuzzy(cleaned_line)

                    # ---- flush buffer and switch section ----
                    if detected_section:
                        # Join buffer content and add to previous section
                        if buffer:
                            buffered_content = " ".join(buffer).strip()
                            if buffered_content:
                                section_map[current_section].append(
                                    buffered_content)
                            buffer = []

                        # Switch to detected section
                        current_section = detected_section
                        LOGGER.debug(f"Detected section: {detected_section}")
                        continue

                    # ---- accumulate content ----
                    buffer.append(cleaned_line)

                    # Flush buffer when reaching a sentence end or max accumulation
                    if (cleaned_line.endswith(('.', ':', '!', '?')) or
                        len(buffer) >= 5 or
                            line_idx == len(raw_lines) - 1):

                        buffered_content = " ".join(buffer).strip()
                        if buffered_content:
                            section_map[current_section].append(
                                buffered_content)
                        buffer = []

                except Exception as line_exc:
                    LOGGER.debug(
                        f"Error processing line {line_idx}: {line_exc}")
                    continue

            # ---- process tables separately ----
            try:
                self._process_tables(tables, section_map)
            except Exception as table_exc:
                print(f"Error processing tables: {table_exc}")
                LOGGER.warning(f"Error processing tables: {table_exc}")

            # ---- format output with logical grouping ----
            final_blocks = []
            processed_sections = set()

            # Define section order for better organization
            section_priority = [
                "summary", "profile", "objective",
                "experience", "work",
                "education", "qualification",
                "skills", "competencies",
                "projects", "academic project",
                "certifications", "certification",
                "achievements", "awards",
                "references", "languages", "interests",
                "personal", "declaration", "general"
            ]

            # Process sections in priority order
            for section_name in section_priority:
                if section_name in section_map and section_map[section_name]:
                    processed_sections.add(section_name)
                    lines = section_map[section_name]

                    try:
                        # Remove duplicates while preserving order
                        unique_lines = []
                        seen = set()
                        for line in lines:
                            if line.strip() and line.strip() not in seen:
                                unique_lines.append(line)
                                seen.add(line.strip())

                        if unique_lines:
                            content = "\n".join(unique_lines)
                            final_blocks.append(
                                self._format_block(section_name, content))
                            LOGGER.info(
                                f"Section '{section_name}': {len(unique_lines)} entries, {len(content)} chars")
                    except Exception as format_exc:
                        print(
                            f"Error formatting section {section_name}: {format_exc}")
                        LOGGER.error(
                            f"Error formatting section {section_name}: {format_exc}")
                        continue

            # Add any remaining unmapped sections
            for section, lines in section_map.items():
                if section not in processed_sections and lines:
                    try:
                        unique_lines = []
                        seen = set()
                        for line in lines:
                            if line.strip() and line.strip() not in seen:
                                unique_lines.append(line)
                                seen.add(line.strip())

                        if unique_lines:
                            content = "\n".join(unique_lines)
                            final_blocks.append(
                                self._format_block(section, content))
                    except Exception as exc:
                        LOGGER.debug(
                            f"Error formatting section {section}: {exc}")

            # Log summary
            LOGGER.info(
                f"Organized into {len(processed_sections)} sections with {sum(len(lines) for lines in section_map.values())} total entries")

            result = "\n".join(final_blocks).strip()
            if not result:
                raise ParsingError("No content extracted after parsing.")

            return result
        except Exception as exc:
            print(f"Error in PDF parse_text: {exc}")
            LOGGER.error(f"Error in PDF parse_text: {exc}")
            raise


class CsvResumeParser(ResumeParserStrategy):
    """Extracts text from CSV resumes by flattening records into lines."""

    def parse_text(self, file_bytes: bytes) -> str:
        try:
            try:
                decoded = file_bytes.decode("utf-8")
            except UnicodeDecodeError as enc_exc:
                print(f"UTF-8 decode failed, trying latin-1: {enc_exc}")
                LOGGER.warning(
                    f"UTF-8 decode failed, trying latin-1: {enc_exc}")
                decoded = file_bytes.decode("latin-1", errors="ignore")

            try:
                reader = csv.DictReader(io.StringIO(decoded))
                rows = list(reader)
                if not rows:
                    raise ParsingError("CSV resume appears empty.")

                lines: list[str] = []
                for row in rows:
                    try:
                        for key, value in row.items():
                            if value is None:
                                continue
                            cleaned = str(value).strip()
                            if cleaned:
                                lines.append(f"{key}: {cleaned}")
                    except Exception as row_exc:
                        print(f"Error processing CSV row: {row_exc}")
                        LOGGER.error(f"Error processing CSV row: {row_exc}")
                        continue

                if not lines:
                    raise ParsingError(
                        "CSV resume has no readable text values.")
                return "\n".join(lines)
            except Exception as csv_exc:
                print(f"Error parsing CSV file: {csv_exc}")
                LOGGER.error(f"Error parsing CSV file: {csv_exc}")
                raise
        except Exception as exc:
            print(f"Error in CSV parse_text: {exc}")
            LOGGER.error(f"Error in CSV parse_text: {exc}")
            raise


class DocxResumeParser(ResumeParserStrategy):
    """Extracts text from DOCX resumes using python-docx, including tables."""

    def parse_text(self, file_bytes: bytes) -> str:
        if DocxDocument is None:
            raise ParsingError("DOCX parsing requires `python-docx`.")

        try:
            document = DocxDocument(io.BytesIO(file_bytes))
            lines: list[str] = []

            # Extract text from paragraphs
            for para in document.paragraphs:
                text = para.text.strip()
                if text:
                    lines.append(text)

            # Extract text from tables
            for table in document.tables:
                for row in table.rows:
                    row_cells: list[str] = []
                    for cell in row.cells:
                        # Each cell may contain multiple paragraphs
                        cell_text = " ".join(
                            p.text.strip() for p in cell.paragraphs if p.text.strip()
                        )
                        if cell_text:
                            row_cells.append(cell_text)
                    # Add the row as a single line with cells separated by " | "
                    if row_cells:
                        lines.append(" | ".join(row_cells))

            text = "\n".join(lines).strip()
            if text:
                return text
        except Exception as exc:
            raise ParsingError(f"Failed to parse DOCX file: {exc}") from exc

        raise ParsingError("DOCX file parsed but no readable text was found.")


class DocResumeParser(ResumeParserStrategy):
    """Extracts text from legacy DOC resumes using textract."""

    def parse_text(self, file_bytes: bytes) -> str:
        if textract is None:
            raise ParsingError(
                "DOC parsing requires `textract` and system parsers (e.g., antiword).")

        temp_path: str | None = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as temp_file:
                temp_file.write(file_bytes)
                temp_path = temp_file.name
            payload = textract.process(temp_path)
            text = payload.decode("utf-8", errors="ignore").strip()
            if text:
                return text
        except Exception as exc:
            raise ParsingError(f"Failed to parse DOC file: {exc}") from exc
        finally:
            if temp_path:
                try:
                    os.remove(temp_path)
                except OSError:
                    LOGGER.warning(
                        "Failed to remove temp DOC file: %s", temp_path)

        raise ParsingError("DOC file parsed but no readable text was found.")
